import argparse
import os
import sys
from typing import List
from PIL import Image
import docx
from io import BytesIO
import uuid
import shutil
import fitz  # PyMuPDF

# Import modul lokal
from qr_utils import (generate_qr, read_qr, analyze_text_encoding, 
                      calculate_qr_capacity, get_optimal_qr_version, 
                      compare_qr_configurations, generate_qr_advanced,
                      generate_secure_qr, read_secure_qr, validate_qr_security)
from lsb_steganography import embed_qr_to_image, extract_qr_from_image

# Import security utilities
import security_utils


def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description='QR Code Watermarking tools menggunakan LSB steganography',
        formatter_class=argparse.RawTextHelpFormatter
    )

    subparsers = parser.add_subparsers(dest='command', help='Perintah yang tersedia')

    # Parser untuk perintah generate_qr
    generate_parser = subparsers.add_parser('generate_qr', help='Generate QR Code')
    generate_parser.add_argument('--data', required=True, help='Data untuk QR Code')
    generate_parser.add_argument('--output', required=True, help='Path untuk menyimpan QR Code')
    generate_parser.add_argument('--version', type=int, help='QR Code version (1-40). Jika tidak diisi, akan otomatis dipilih')
    generate_parser.add_argument('--error-correction', choices=['L', 'M', 'Q', 'H'], 
                                default='M', help='Tingkat koreksi error (default: M)')
    generate_parser.add_argument('--box-size', type=int, default=10, 
                                help='Ukuran setiap kotak dalam piksel (default: 10)')
    generate_parser.add_argument('--border', type=int, default=4,
                                help='Lebar border dalam kotak (default: 4)')
    generate_parser.add_argument('--analyze', action='store_true',
                                help='Tampilkan analisis konfigurasi QR Code optimal')
    # Security-related arguments
    generate_parser.add_argument('--secure', action='store_true',
                                help='Generate secure QR with encryption')
    generate_parser.add_argument('--document-path', type=str,
                                help='Path to document for security key generation (required with --secure)')

    # Parser untuk perintah embed_docx
    embed_parser = subparsers.add_parser('embed_docx', help='Embed QR Code watermark ke dokumen')
    embed_parser.add_argument('--docx', required=True, help='Path ke dokumen .docx')
    embed_parser.add_argument('--qr', required=True, help='Path ke QR Code watermark')
    embed_parser.add_argument('--output', required=True, help='Path untuk menyimpan dokumen hasil')

    # Parser untuk perintah embed_pdf
    embed_pdf_parser = subparsers.add_parser('embed_pdf', help='Embed QR Code watermark ke dokumen PDF')
    embed_pdf_parser.add_argument('--pdf', required=True, help='Path ke dokumen .pdf')
    embed_pdf_parser.add_argument('--qr', required=True, help='Path ke QR Code watermark')
    embed_pdf_parser.add_argument('--output', required=True, help='Path untuk menyimpan dokumen hasil')

    # Parser untuk perintah extract_docx
    extract_parser = subparsers.add_parser('extract_docx', help='Extract QR Code watermark dari dokumen')
    extract_parser.add_argument('--docx', required=True, help='Path ke dokumen .docx')
    extract_parser.add_argument('--output_dir', required=True, help='Direktori untuk menyimpan QR Code hasil ekstraksi')

    # Parser untuk perintah extract_pdf
    extract_pdf_parser = subparsers.add_parser('extract_pdf', help='Extract QR Code watermark dari dokumen PDF')
    extract_pdf_parser.add_argument('--pdf', required=True, help='Path ke dokumen .pdf')
    extract_pdf_parser.add_argument('--output_dir', required=True, help='Direktori untuk menyimpan QR Code hasil ekstraksi')

    # Parser untuk perintah generate_key (NEW)
    generate_key_parser = subparsers.add_parser('generate_key', help='Generate security key for document')
    generate_key_parser.add_argument('--document', required=True, help='Path to document file')
    generate_key_parser.add_argument('--additional-data', type=str, default='',
                                    help='Additional data to include in key generation (optional)')
    generate_key_parser.add_argument('--save-key', type=str,
                                    help='Path to save the generated key to file (optional)')

    # Parser untuk perintah validate_security (NEW)
    validate_security_parser = subparsers.add_parser('validate_security', help='Validate QR-document security pairing')
    validate_security_parser.add_argument('--qr-image', required=True, help='Path to QR code image')
    validate_security_parser.add_argument('--document', required=True, help='Path to document file')
    validate_security_parser.add_argument('--key', type=str,
                                         help='Document security key (if not provided, will generate from document)')
    validate_security_parser.add_argument('--detailed', action='store_true',
                                         help='Show detailed validation report')

    return parser.parse_args()


def analyze_qr_options(data: str, target_image_sizes: list = None) -> dict:
    """
    Analyze and recommend optimal QR code settings for the given data.
    
    Args:
        data (str): The data to encode in the QR code
        target_image_sizes (list, optional): List of target image sizes [(width, height), ...]
                                           for steganography capacity validation
    
    Returns:
        dict: Analysis results with recommendations
    """
    if not data:
        raise ValueError("Data cannot be empty")
    
    try:
        # Analyze encoding
        optimal_encoding = analyze_text_encoding(data)
        data_length = len(data)
        
        print(f"[*] Analyzing QR options for data: '{data[:50]}{'...' if len(data) > 50 else ''}'")
        print(f"[*] Data length: {data_length} characters")
        print(f"[*] Optimal encoding: {optimal_encoding}")
        
        # Get optimal versions for different error correction levels
        optimal_versions = {}
        for error_level in ['L', 'M', 'Q', 'H']:
            try:
                version = get_optimal_qr_version(data_length, optimal_encoding, error_level)
                optimal_versions[error_level] = version
            except ValueError as e:
                optimal_versions[error_level] = f"Error: {str(e)}"
        
        # Compare configurations
        comparison_versions = [v for v in optimal_versions.values() if isinstance(v, int)]
        if comparison_versions:
            comparison_versions = sorted(set(comparison_versions))[:3]  # Top 3 versions
            comparison = compare_qr_configurations(data, comparison_versions)
        else:
            comparison = None
        
        # Validate against target image sizes if provided
        steganography_validation = None
        if target_image_sizes:
            steganography_validation = []
            for width, height in target_image_sizes:
                # Estimate image capacity for steganography (rough calculation)
                # Typically, we can hide 1 bit per pixel in LSB steganography
                image_capacity_bits = width * height
                image_capacity_bytes = image_capacity_bits // 8
                
                # QR code size estimation (very rough)
                for error_level, version in optimal_versions.items():
                    if isinstance(version, int):
                        # QR modules = (4 * version + 17)^2
                        qr_modules = (4 * version + 17) ** 2
                        # Assume each module needs at least 1 byte for storage
                        estimated_qr_bytes = qr_modules
                        
                        can_fit = image_capacity_bytes >= estimated_qr_bytes
                        steganography_validation.append({
                            'image_size': f"{width}x{height}",
                            'error_level': error_level,
                            'version': version,
                            'can_fit': can_fit,
                            'image_capacity': image_capacity_bytes,
                            'estimated_qr_size': estimated_qr_bytes
                        })
        
        # Prepare recommendation
        recommended_config = None
        if 'M' in optimal_versions and isinstance(optimal_versions['M'], int):
            recommended_config = {
                'version': optimal_versions['M'],
                'error_correction': 'M',
                'encoding': optimal_encoding,
                'reason': 'Balanced error correction and size'
            }
        elif any(isinstance(v, int) for v in optimal_versions.values()):
            # Find the first working version
            for level in ['L', 'Q', 'H']:
                if isinstance(optimal_versions.get(level), int):
                    recommended_config = {
                        'version': optimal_versions[level],
                        'error_correction': level,
                        'encoding': optimal_encoding,
                        'reason': f'Best available option with {level} error correction'
                    }
                    break
        
        result = {
            'data_info': {
                'length': data_length,
                'optimal_encoding': optimal_encoding
            },
            'optimal_versions': optimal_versions,
            'comparison': comparison,
            'steganography_validation': steganography_validation,
            'recommended_config': recommended_config
        }
        
        # Print recommendations
        print("\n[*] Analysis Results:")
        print(f"    Optimal encoding: {optimal_encoding}")
        print("    Minimum versions needed:")
        for level, version in optimal_versions.items():
            print(f"      Error level {level}: {version}")
        
        if recommended_config:
            print(f"\n[*] Recommended configuration:")
            print(f"    Version: {recommended_config['version']}")
            print(f"    Error correction: {recommended_config['error_correction']}")
            print(f"    Reason: {recommended_config['reason']}")
        
        return result
        
    except Exception as e:
        print(f"[!] Error during QR analysis: {str(e)}")
        raise


def generate_qr_code(data: str, output_path: str, version: int = None, 
                    error_correction: str = 'M', box_size: int = 10, 
                    border: int = 4, analyze: bool = False, secure: bool = False,
                    document_path: str = None) -> bool:
    """
    Generate QR Code with advanced configuration options and save it to the specified path.

    Args:
        data (str): Text data to encode in the QR Code
        output_path (str): Path to save the generated QR Code
        version (int, optional): QR version (1-40). If None, automatically determined
        error_correction (str): Error correction level ('L', 'M', 'Q', 'H')
        box_size (int): Size of each box in pixels
        border (int): Border size in boxes
        analyze (bool): Whether to show analysis of QR configuration
        secure (bool): Whether to generate secure encrypted QR code
        document_path (str): Path to document for security key generation (required if secure=True)

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Validate secure mode requirements
        if secure:
            if not document_path:
                raise ValueError("Document path is required for secure QR generation")
            if not os.path.exists(document_path):
                raise FileNotFoundError(f"Document not found: {document_path}")
            
            print(f"[*] Generating secure QR code with document-based encryption...")
            
            # Generate document key
            document_key = security_utils.generate_document_key(document_path)
            print(f"[*] Document security key generated")
            
            # Generate secure QR code
            img = generate_secure_qr(data, document_key, output_path)
            
            print(f"[*] Secure QR Code berhasil dibuat dengan data: '{data}'")
            print(f"[*] Tersimpan di: {output_path}")
            print(f"[*] Document key: {document_key}")
            print(f"[!] IMPORTANT: Save the document key to decrypt QR data later!")
            return True
        
        # Non-secure mode (original functionality)
        # Perform analysis if requested
        if analyze:
            try:
                analyze_qr_options(data)
                print()  # Add spacing before generation
            except Exception as e:
                print(f"[!] Warning: Analysis failed: {str(e)}")
        
        # Validate QR capacity before generation
        if version is not None:
            try:
                encoding = analyze_text_encoding(data)
                capacity = calculate_qr_capacity(version, error_correction, encoding)
                if len(data) > capacity:
                    print(f"[!] Warning: Data length ({len(data)}) exceeds QR capacity ({capacity}) for version {version} with {error_correction} error correction")
                    print(f"[*] Trying to find optimal version...")
                    version = None  # Let the advanced generator find optimal version
            except Exception as e:
                print(f"[!] Warning: Capacity validation failed: {str(e)}")
        
        # Use advanced QR generation
        img = generate_qr_advanced(
            data=data,
            version=version,
            error_correction=error_correction,
            box_size=box_size,
            border=border,
            output_path=output_path
        )
        
        print(f"[*] QR Code berhasil dibuat dengan data: '{data}'")
        print(f"[*] Tersimpan di: {output_path}")
        return True
        
    except Exception as e:
        print(f"[!] Error saat membuat QR code: {str(e)}")
        # Fallback to basic generation (only for non-secure mode)
        if not secure:
            try:
                print("[*] Mencoba dengan metode dasar...")
                generate_qr(data, output_path)
                print(f"[*] QR Code berhasil dibuat dengan metode dasar")
                return True
            except Exception as fallback_e:
                print(f"[!] Metode dasar juga gagal: {str(fallback_e)}")
        return False


def extract_images_from_docx(docx_path: str, output_dir: str) -> List[str]:
    """
    Extract all images from a docx file and save them to output directory.

    Args:
        docx_path: Path to the .docx file
        output_dir: Directory to save extracted images

    Returns:
        List[str]: List of paths to the extracted images
    """
    try:
        # Make sure the output directory exists
        os.makedirs(output_dir, exist_ok=True)

        # Open the document
        doc = docx.Document(docx_path)

        # Initialize an empty list for image paths
        image_paths = []

        # Loop through all the document parts that could contain images
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                # Get the image data
                image_data = rel.target_part.blob

                # Create a file name for the image
                image_filename = f"image_{image_count}.png"
                image_path = os.path.join(output_dir, image_filename)

                # Save the image
                with open(image_path, "wb") as f:
                    f.write(image_data)

                image_paths.append(image_path)
                image_count += 1

        if image_count > 0:
            print(f"[*] Berhasil mengekstrak {image_count} gambar dari dokumen")
        else:
            print("[!] Tidak ada gambar yang ditemukan dalam dokumen")

        return image_paths
    except Exception as e:
        print(f"[!] Error saat mengekstrak gambar dari dokumen: {str(e)}")
        return []


def extract_images_from_pdf(pdf_path: str, output_dir: str) -> List[str]:
    """
    Extract all images from a PDF file and save them to output directory.

    Args:
        pdf_path: Path to the .pdf file
        output_dir: Directory to save extracted images

    Returns:
        List[str]: List of paths to the extracted images
    """
    try:
        # Make sure the output directory exists
        os.makedirs(output_dir, exist_ok=True)

        # Open the PDF document
        doc = fitz.open(pdf_path)

        # Initialize an empty list for image paths
        image_paths = []
        image_count = 0

        # Loop through all pages
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images(full=True)

            # Loop through images on this page
            for img_index, img in enumerate(image_list):
                # Get the XREF of the image
                xref = img[0]
                
                # Extract the image bytes
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Convert to PNG if not already PNG
                if image_ext.lower() != "png":
                    # Open with PIL and convert to PNG
                    image = Image.open(BytesIO(image_bytes))
                    image_filename = f"image_{image_count}.png"
                    image_path = os.path.join(output_dir, image_filename)
                    image.save(image_path, "PNG")
                else:
                    # Save directly as PNG
                    image_filename = f"image_{image_count}.png"
                    image_path = os.path.join(output_dir, image_filename)
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)

                image_paths.append(image_path)
                image_count += 1

        if image_count > 0:
            print(f"[*] Berhasil mengekstrak {image_count} gambar dari PDF")
        else:
            print("[!] Tidak ada gambar yang ditemukan dalam PDF")

        doc.close()
        return image_paths
    except Exception as e:
        print(f"[!] Error saat mengekstrak gambar dari PDF: {str(e)}")
        return []


def replace_images_in_docx(docx_path: str, original_images: List[str],
                          watermarked_images: List[str], output_path: str) -> bool:
    """
    Replace images in a docx document with watermarked versions.

    Args:
        docx_path: Path to the original .docx file
        original_images: List of paths to the original images
        watermarked_images: List of paths to the watermarked images
        output_path: Path to save the new document

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Create a mapping between original and watermarked images
        image_map = {}
        for orig, watermarked in zip(original_images, watermarked_images):
            image_map[os.path.basename(orig)] = watermarked

        # Make a copy of the document
        doc = docx.Document(docx_path)

        # Loop through all the document parts that could contain images
        images_replaced = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                rel_id = rel.rId
                target_ref = rel.target_ref
                image_name = target_ref.split("/")[-1]

                # Check if we have a watermarked version of this image
                base_image_name = f"image_{images_replaced}.png"
                if base_image_name in [os.path.basename(img) for img in image_map.keys()]:
                    watermarked_path = image_map[base_image_name]

                    # Read the watermarked image
                    with open(watermarked_path, "rb") as f:
                        watermarked_data = f.read()

                    # Replace the image in the document
                    # Since we can't set blob directly, we need to modify the part differently
                    try:
                        # Method 1: Try to modify part's _blob attribute (not recommended but works in some versions)
                        if hasattr(rel.target_part, '_blob'):
                            rel.target_part._blob = watermarked_data
                            images_replaced += 1
                        # Method 2: Create a new part and replace the old one
                        else:
                            # Get the content type of the image
                            content_type = rel.target_part.content_type
                            
                            # Create a new image part with the watermarked data
                            new_img_part = doc.part.package.blob_storage.new_part(
                                content_type, 
                                doc.part.package.next_partname(f"/word/media/image{uuid.uuid4().hex}.png")
                            )
                            
                            # Write the watermarked data to the new part
                            with new_img_part.open('wb') as f:
                                f.write(watermarked_data)
                                
                            # Update the relationship to point to the new part
                            doc.part.rels[rel_id].target_part = new_img_part
                            images_replaced += 1
                    except Exception as img_e:
                        print(f"[!] Warning: Error mengganti gambar: {str(img_e)}")
                        # Try alternative method
                        try:
                            # Method 3: Create a new relationship and delete the old one
                            # Get original part name for reference
                            original_part_name = rel.target_part.partname
                            
                            # Add a new image part with the watermarked data
                            new_rid = doc.part.add_image(watermarked_data)
                            
                            # Find all references to the old rId and replace with new_rid
                            for element in doc.element.body.iter():
                                if hasattr(element, 'attrib'):
                                    for key, value in element.attrib.items():
                                        if value == rel_id:
                                            element.attrib[key] = new_rid
                            
                            images_replaced += 1
                        except Exception as alt_e:
                            print(f"[!] Warning: Alternative method juga gagal: {str(alt_e)}")

        # Save the document
        doc.save(output_path)

        if images_replaced > 0:
            print(f"[*] Berhasil mengganti {images_replaced} gambar dengan versi watermark")
            print(f"[*] Dokumen tersimpan di: {output_path}")
            return True
        else:
            print("[!] Tidak ada gambar yang diganti dalam dokumen")
            return False
    except Exception as e:
        print(f"[!] Error saat mengganti gambar dalam dokumen: {str(e)}")
        return False


def replace_images_in_pdf(pdf_path: str, original_images: List[str],
                         watermarked_images: List[str], output_path: str) -> bool:
    """
    Replace images in a PDF document with watermarked versions.

    Args:
        pdf_path: Path to the original .pdf file
        original_images: List of paths to the original images
        watermarked_images: List of paths to the watermarked images
        output_path: Path to save the new document

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Open the PDF document
        doc = fitz.open(pdf_path)

        images_replaced = 0
        image_index = 0

        # Loop through all pages
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images(full=True)

            # Loop through images on this page
            for img_index, img in enumerate(image_list):
                if image_index < len(watermarked_images):
                    # Get the XREF of the image
                    xref = img[0]
                    
                    # Read the watermarked image
                    watermarked_path = watermarked_images[image_index]
                    
                    # Read watermarked image data
                    with open(watermarked_path, "rb") as f:
                        watermarked_data = f.read()
                    
                    # Replace the image in the PDF
                    try:
                        # Get image rectangle on the page
                        img_rects = page.get_image_rects(xref)
                        
                        if img_rects:
                            # Get the first rectangle (in case there are multiple instances)
                            rect = img_rects[0]
                            
                            # Delete the old image
                            page.delete_image(xref)
                            
                            # Insert the new watermarked image at the same position
                            page.insert_image(rect, stream=watermarked_data)
                            
                            images_replaced += 1
                            print(f"[*] Mengganti gambar {image_index + 1} pada halaman {page_num + 1}")
                        else:
                            print(f"[!] Warning: Tidak dapat menemukan posisi gambar {image_index + 1}")
                    
                    except Exception as img_e:
                        print(f"[!] Warning: Error mengganti gambar {image_index + 1}: {str(img_e)}")
                        
                    image_index += 1

        # Save the document
        doc.save(output_path)
        doc.close()

        if images_replaced > 0:
            print(f"[*] Berhasil mengganti {images_replaced} gambar dengan versi watermark")
            print(f"[*] Dokumen tersimpan di: {output_path}")
            return True
        else:
            print("[!] Tidak ada gambar yang diganti dalam dokumen")
            return False
            
    except Exception as e:
        print(f"[!] Error saat mengganti gambar dalam PDF: {str(e)}")
        return False


def analyze_docx_images(docx_path: str) -> dict:
    """
    Analyze images in a DOCX document to get size information for QR optimization.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        dict: Analysis results with image dimensions
    """
    try:
        import io
        try:
            from docx import Document
        except ImportError:
            print("[!] python-docx library not installed, skipping DOCX image analysis")
            return {'success': False, 'error': 'python-docx not installed', 'images': []}
        
        doc = Document(docx_path)
        images_info = []
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_data))
                    width, height = img.size
                    images_info.append({
                        'width': width,
                        'height': height,
                        'format': img.format,
                        'mode': img.mode
                    })
                except Exception as e:
                    print(f"[!] Could not analyze image: {e}")
                    continue
        
        return {
            'success': True,
            'images': images_info,
            'total_images': len(images_info)
        }
        
    except Exception as e:
        print(f"[!] Error analyzing DOCX images: {e}")
        return {'success': False, 'error': str(e), 'images': []}


def analyze_pdf_images(pdf_path: str) -> dict:
    """
    Analyze images in a PDF document to get size information for QR optimization.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        dict: Analysis results with image dimensions
    """
    try:
        import io
        try:
            import fitz  # PyMuPDF
        except ImportError:
            print("[!] PyMuPDF library not installed, skipping PDF image analysis")
            return {'success': False, 'error': 'PyMuPDF not installed', 'images': []}
        
        pdf_doc = fitz.open(pdf_path)
        images_info = []
        
        for page_num in range(pdf_doc.page_count):
            page = pdf_doc[page_num]
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = pdf_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    img_pil = Image.open(io.BytesIO(image_bytes))
                    width, height = img_pil.size
                    
                    images_info.append({
                        'width': width,
                        'height': height,
                        'format': img_pil.format,
                        'mode': img_pil.mode,
                        'page': page_num + 1
                    })
                except Exception as e:
                    print(f"[!] Could not analyze PDF image: {e}")
                    continue
        
        pdf_doc.close()
        
        return {
            'success': True,
            'images': images_info,
            'total_images': len(images_info)
        }
        
    except Exception as e:
        print(f"[!] Error analyzing PDF images: {e}")
        return {'success': False, 'error': str(e), 'images': []}


# ====== NEW SECURITY FUNCTIONS ======

def generate_document_key_command(document_path: str, additional_data: str = "", save_key_path: str = None) -> bool:
    """
    CLI command to generate and display document security key.
    
    Args:
        document_path (str): Path to the document file
        additional_data (str): Optional additional data for key generation
        save_key_path (str): Optional path to save the key to file
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        if not os.path.exists(document_path):
            print(f"[!] Error: Document not found: {document_path}")
            return False
        
        print(f"[*] Generating security key for document: {os.path.basename(document_path)}")
        
        # Generate document key
        document_key = security_utils.generate_document_key(document_path, additional_data)
        
        # Generate document hash for validation
        document_hash = security_utils.generate_document_hash(document_path)
        
        # Create key signature for validation
        key_signature = security_utils.create_key_signature(document_key, document_hash)
        
        print(f"\n[*] Document Security Information:")
        print(f"    Document: {os.path.basename(document_path)}")
        print(f"    Document Hash: {document_hash}")
        print(f"    Security Key: {document_key}")
        print(f"    Key Signature: {key_signature}")
        
        # Save key to file if requested
        if save_key_path:
            try:
                key_info = {
                    'document_path': document_path,
                    'document_name': os.path.basename(document_path),
                    'document_hash': document_hash,
                    'security_key': document_key,
                    'key_signature': key_signature,
                    'additional_data': additional_data,
                    'generated_at': security_utils.datetime.now().isoformat()
                }
                
                import json
                with open(save_key_path, 'w') as f:
                    json.dump(key_info, f, indent=2)
                
                print(f"[*] Key information saved to: {save_key_path}")
                
            except Exception as e:
                print(f"[!] Warning: Could not save key to file: {e}")
        
        print(f"\n[!] IMPORTANT: Save the security key securely!")
        print(f"[!] You need this key to decrypt QR codes generated for this document.")
        
        return True
        
    except Exception as e:
        print(f"[!] Error generating document key: {e}")
        return False


def validate_qr_document_pair(qr_image_path: str, document_path: str, document_key: str = None, detailed: bool = False) -> bool:
    """
    CLI command to validate if QR belongs to specific document.
    
    Args:
        qr_image_path (str): Path to QR code image
        document_path (str): Path to document file
        document_key (str): Document security key (optional, will generate if not provided)
        detailed (bool): Whether to show detailed validation report
        
    Returns:
        bool: True if validation passes, False otherwise
    """
    try:
        # Validate inputs
        if not os.path.exists(qr_image_path):
            print(f"[!] Error: QR image not found: {qr_image_path}")
            return False
        
        if not os.path.exists(document_path):
            print(f"[!] Error: Document not found: {document_path}")
            return False
        
        print(f"[*] Validating QR-document security pairing...")
        print(f"    QR Image: {os.path.basename(qr_image_path)}")
        print(f"    Document: {os.path.basename(document_path)}")
        
        # Generate document key if not provided
        if not document_key:
            print(f"[*] Generating document key...")
            document_key = security_utils.generate_document_key(document_path)
            print(f"[*] Document key generated")
        else:
            print(f"[*] Using provided document key")
        
        # Read QR code data
        try:
            qr_data_list = read_qr(qr_image_path)
            if not qr_data_list:
                print(f"[!] Error: No QR code data found in image")
                return False
            
            qr_data = qr_data_list[0]  # Use first QR code
            print(f"[*] QR data read successfully ({len(qr_data)} chars)")
            
        except Exception as e:
            print(f"[!] Error reading QR code: {e}")
            return False
        
        # Generate document hash for validation
        document_hash = security_utils.generate_document_hash(document_path)
        
        # Perform comprehensive security validation
        if detailed:
            print(f"\n[*] Performing detailed security validation...")
            validation_results = validate_qr_security(qr_data, document_key, document_hash)
            
            print(f"\n[*] Detailed Validation Results:")
            print(f"    Overall Valid: {validation_results['overall_valid']}")
            print(f"    Security Score: {validation_results.get('security_score', 0):.2f}/1.0")
            print(f"    Validation Timestamp: {validation_results['timestamp']}")
            
            print(f"\n[*] Individual Checks:")
            for check, result in validation_results.get('validations', {}).items():
                status = "✓ PASS" if result else "✗ FAIL" if result is False else "- SKIP"
                print(f"      {check}: {status}")
            
            if validation_results.get('warnings'):
                print(f"\n[!] Warnings:")
                for warning in validation_results['warnings']:
                    print(f"      - {warning}")
            
            if validation_results.get('errors'):
                print(f"\n[!] Errors:")
                for error in validation_results['errors']:
                    print(f"      - {error}")
            
            # Try to decrypt QR data
            try:
                decrypted_data = read_secure_qr(qr_image_path, document_key)
                print(f"\n[*] QR Data Decryption: SUCCESS")
                print(f"    Decrypted content: '{decrypted_data}'")
            except Exception as e:
                print(f"\n[!] QR Data Decryption: FAILED")
                print(f"    Error: {e}")
            
            return validation_results['overall_valid']
        
        else:
            # Simple validation using security utils
            try:
                is_authorized = security_utils.is_qr_authorized_for_document(qr_data, document_key, document_path)
                
                if is_authorized:
                    print(f"\n[*] ✅ QR-Document Validation: PASSED")
                    print(f"[*] The QR code is authorized for this document")
                    
                    # Try to decrypt and show content
                    try:
                        decrypted_data = read_secure_qr(qr_image_path, document_key)
                        print(f"[*] Decrypted QR content: '{decrypted_data}'")
                    except Exception as e:
                        print(f"[!] Could not decrypt QR content: {e}")
                else:
                    print(f"\n[!] ❌ QR-Document Validation: FAILED")
                    print(f"[!] The QR code is NOT authorized for this document")
                
                return is_authorized
                
            except Exception as e:
                print(f"[!] Validation error: {e}")
                return False
    
    except Exception as e:
        print(f"[!] Error during validation: {e}")
        return False


def embed_watermark_to_docx(docx_path: str, qr_path: str, output_path: str, validate_security: bool = False, document_key: str = None) -> dict:
    """
    Process a docx document to add QR watermarks to all images.

    Args:
        docx_path: Path to the original .docx file
        qr_path: Path to the QR code image to embed
        output_path: Path to save the watermarked document
        validate_security: Whether to perform security validation before embedding
        document_key: Document security key for validation (optional)

    Returns:
        dict: Result dictionary with success status and processed image info
    """
    try:
        # Security validation if requested
        if validate_security:
            print(f"[*] Performing security validation before embedding...")
            
            # Generate document key if not provided
            if not document_key:
                document_key = security_utils.generate_document_key(docx_path)
            
            # Validate QR-document pairing
            try:
                qr_data_list = read_qr(qr_path)
                if qr_data_list:
                    qr_data = qr_data_list[0]
                    document_hash = security_utils.generate_document_hash(docx_path)
                    
                    # Use comprehensive validation
                    validation_results = validate_qr_security(qr_data, document_key, document_hash)
                    
                    if not validation_results['overall_valid']:
                        print(f"[!] Security validation failed - QR code is not authorized for this document")
                        return {
                            'success': False,
                            'error': 'Security validation failed',
                            'validation_results': validation_results
                        }
                    else:
                        print(f"[*] ✅ Security validation passed - proceeding with embedding")
                else:
                    print(f"[!] Warning: Could not read QR code for validation")
            except Exception as e:
                print(f"[!] Warning: Security validation error: {e}")
        
        # Validate QR code for steganography capacity
        try:
            qr_img = Image.open(qr_path)
            qr_width, qr_height = qr_img.size
            qr_pixels = qr_width * qr_height
            print(f"[*] QR Code size: {qr_width}x{qr_height} ({qr_pixels} pixels)")
        except Exception as e:
            print(f"[!] Warning: Could not analyze QR code: {str(e)}")
            qr_width, qr_height, qr_pixels = 21, 21, 441  # Default values

        # Create a unique temporary directory to store extracted and watermarked images
        temp_dir_name = f"temp_embed_{uuid.uuid4().hex}"
        temp_dir = os.path.join(os.path.dirname(output_path), temp_dir_name)
        os.makedirs(temp_dir, exist_ok=True)

        # Create a public directory to store images for viewing
        public_dir_name = f"processed_{uuid.uuid4().hex}"
        public_dir = os.path.join(os.path.dirname(output_path), public_dir_name)
        os.makedirs(public_dir, exist_ok=True)

        # Extract images from the document
        print(f"[*] Mengekstrak gambar dari dokumen: {docx_path}")
        extracted_images = extract_images_from_docx(docx_path, temp_dir)

        if not extracted_images:
            print("[!] Dokumen ini tidak mengandung gambar")
            # Clean up temporary directory
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except Exception as e:
                    print(f"[!] Warning: Tidak dapat menghapus direktori temp: {str(e)}")
            # Return False with a specific message that can be detected by the app
            raise ValueError("NO_IMAGES_FOUND")

        # Validate each image for steganography capacity
        valid_images = []
        for img_path in extracted_images:
            try:
                img = Image.open(img_path)
                img_width, img_height = img.size
                img_pixels = img_width * img_height
                
                # Rough capacity check: image should have at least 4x the QR pixels for good steganography
                # This is a conservative estimate to ensure quality
                min_required_pixels = qr_pixels * 4
                
                if img_pixels >= min_required_pixels:
                    valid_images.append((img_path, img_width, img_height))
                    print(f"[*] Image {os.path.basename(img_path)}: {img_width}x{img_height} - OK for steganography")
                else:
                    print(f"[!] Warning: Image {os.path.basename(img_path)}: {img_width}x{img_height} - Mungkin terlalu kecil untuk QR {qr_width}x{qr_height}")
                    # Still include it, but with warning
                    valid_images.append((img_path, img_width, img_height))
            except Exception as e:
                print(f"[!] Warning: Could not analyze image {img_path}: {str(e)}")
                valid_images.append((img_path, 0, 0))  # Include with unknown size

        # Prepare to store info about processed images
        processed_images = []

        # Watermark each image
        watermarked_images = []
        for i, (img_path, img_w, img_h) in enumerate(valid_images):
            print(f"[*] Menyisipkan watermark QR Code ke gambar {i + 1}/{len(valid_images)}")
            watermarked_path = os.path.join(temp_dir, f"watermarked_{i}.png")
            
            # Create public copies for display
            original_public_name = f"original_{i}.png"
            watermarked_public_name = f"watermarked_{i}.png"
            original_public_path = os.path.join(public_dir, original_public_name)
            watermarked_public_path = os.path.join(public_dir, watermarked_public_name)
            
            # Copy original to public directory
            shutil.copy(img_path, original_public_path)
            
            try:
                # Perform the watermarking
                embed_qr_to_image(img_path, qr_path, watermarked_path, resize_qr_if_needed=True)
                watermarked_images.append(watermarked_path)
                
                # Copy watermarked image to public directory
                shutil.copy(watermarked_path, watermarked_public_path)
                
                # Store info about this image pair with capacity info
                processed_images.append({
                    "index": i,
                    "original": f"{public_dir_name}/{original_public_name}",
                    "watermarked": f"{public_dir_name}/{watermarked_public_name}",
                    "image_size": f"{img_w}x{img_h}",
                    "capacity_ok": img_w * img_h >= qr_pixels * 4 if img_w > 0 and img_h > 0 else None
                })
                
            except Exception as e:
                print(f"[!] Gagal watermark gambar {img_path}: {str(e)}")
                # Continue with other images even if one fails

        # Replace images in the document with watermarked versions
        print(f"[*] Mengganti gambar dalam dokumen dengan versi watermark")
        success = replace_images_in_docx(docx_path, [path for path, _, _ in valid_images], watermarked_images, output_path)

        # Clean up temporary files
        for path, _, _ in valid_images:
            if os.path.exists(path) and not os.path.basename(path).startswith("original_") and not os.path.basename(path).startswith("watermarked_"):
                os.remove(path)
        for path in watermarked_images:
            if os.path.exists(path) and not os.path.basename(path).startswith("original_") and not os.path.basename(path).startswith("watermarked_"):
                os.remove(path)

        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                print(f"[!] Warning: Tidak dapat menghapus direktori temp: {str(e)}")

        # Copy QR code to public directory for display
        qr_public_name = "watermark_qr.png"
        qr_public_path = os.path.join(public_dir, qr_public_name)
        shutil.copy(qr_path, qr_public_path)

        # Get QR code dimensions
        try:
            qr_img = Image.open(qr_path)
            qr_width, qr_height = qr_img.size
            qr_info = {
                "width": qr_width,
                "height": qr_height
            }
        except Exception as e:
            print(f"[!] Warning: Couldn't get QR dimensions: {str(e)}")
            qr_info = {"width": 21, "height": 21}  # Default QR size

        # Return success status and image info
        return {
            "success": success, 
            "processed_images": processed_images,
            "qr_image": f"{public_dir_name}/{qr_public_name}",
            "public_dir": public_dir_name,
            "qr_info": qr_info
        }
        
    except ValueError as ve:
        if str(ve) == "NO_IMAGES_FOUND":
            # Propagate the specific error
            raise ve
        print(f"[!] Error saat proses watermarking dokumen: {str(ve)}")
        return {"success": False, "error": str(ve)}
    except Exception as e:
        print(f"[!] Error saat proses watermarking dokumen: {str(e)}")
        return {"success": False, "error": str(e)}


def embed_watermark_to_pdf(pdf_path: str, qr_path: str, output_path: str, validate_security: bool = False, document_key: str = None) -> dict:
    """
    Process a PDF document to add QR watermarks to all images.

    Args:
        pdf_path: Path to the original .pdf file
        qr_path: Path to the QR code image to embed
        output_path: Path to save the watermarked document
        validate_security: Whether to perform security validation before embedding
        document_key: Document security key for validation (optional)

    Returns:
        dict: Result dictionary with success status and processed image info
    """
    try:
        # Security validation if requested
        if validate_security:
            print(f"[*] Performing security validation before embedding...")
            
            # Generate document key if not provided
            if not document_key:
                document_key = security_utils.generate_document_key(pdf_path)
            
            # Validate QR-document pairing
            try:
                qr_data_list = read_qr(qr_path)
                if qr_data_list:
                    qr_data = qr_data_list[0]
                    document_hash = security_utils.generate_document_hash(pdf_path)
                    
                    # Use comprehensive validation
                    validation_results = validate_qr_security(qr_data, document_key, document_hash)
                    
                    if not validation_results['overall_valid']:
                        print(f"[!] Security validation failed - QR code is not authorized for this document")
                        return {
                            'success': False,
                            'error': 'Security validation failed',
                            'validation_results': validation_results
                        }
                    else:
                        print(f"[*] ✅ Security validation passed - proceeding with embedding")
                else:
                    print(f"[!] Warning: Could not read QR code for validation")
            except Exception as e:
                print(f"[!] Warning: Security validation error: {e}")
        
        # Validate QR code for steganography capacity
        try:
            qr_img = Image.open(qr_path)
            qr_width, qr_height = qr_img.size
            qr_pixels = qr_width * qr_height
            print(f"[*] QR Code size: {qr_width}x{qr_height} ({qr_pixels} pixels)")
        except Exception as e:
            print(f"[!] Warning: Could not analyze QR code: {str(e)}")
            qr_width, qr_height, qr_pixels = 21, 21, 441  # Default values

        # Create a unique temporary directory to store extracted and watermarked images
        temp_dir_name = f"temp_embed_{uuid.uuid4().hex}"
        temp_dir = os.path.join(os.path.dirname(output_path), temp_dir_name)
        os.makedirs(temp_dir, exist_ok=True)

        # Create a public directory to store images for viewing
        public_dir_name = f"processed_{uuid.uuid4().hex}"
        public_dir = os.path.join(os.path.dirname(output_path), public_dir_name)
        os.makedirs(public_dir, exist_ok=True)

        # Extract images from the document
        print(f"[*] Mengekstrak gambar dari PDF: {pdf_path}")
        extracted_images = extract_images_from_pdf(pdf_path, temp_dir)

        if not extracted_images:
            print("[!] PDF ini tidak mengandung gambar")
            # Clean up temporary directory
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except Exception as e:
                    print(f"[!] Warning: Tidak dapat menghapus direktori temp: {str(e)}")
            # Return False with a specific message that can be detected by the app
            raise ValueError("NO_IMAGES_FOUND")

        # Validate each image for steganography capacity
        valid_images = []
        for img_path in extracted_images:
            try:
                img = Image.open(img_path)
                img_width, img_height = img.size
                img_pixels = img_width * img_height
                
                # Rough capacity check: image should have at least 4x the QR pixels for good steganography
                min_required_pixels = qr_pixels * 4
                
                if img_pixels >= min_required_pixels:
                    valid_images.append((img_path, img_width, img_height))
                    print(f"[*] Image {os.path.basename(img_path)}: {img_width}x{img_height} - OK for steganography")
                else:
                    print(f"[!] Warning: Image {os.path.basename(img_path)}: {img_width}x{img_height} - Mungkin terlalu kecil untuk QR {qr_width}x{qr_height}")
                    # Still include it, but with warning
                    valid_images.append((img_path, img_width, img_height))
            except Exception as e:
                print(f"[!] Warning: Could not analyze image {img_path}: {str(e)}")
                valid_images.append((img_path, 0, 0))  # Include with unknown size

        # Prepare to store info about processed images
        processed_images = []

        # Watermark each image
        watermarked_images = []
        for i, (img_path, img_w, img_h) in enumerate(valid_images):
            print(f"[*] Menyisipkan watermark QR Code ke gambar {i + 1}/{len(valid_images)}")
            watermarked_path = os.path.join(temp_dir, f"watermarked_{i}.png")
            
            # Create public copies for display
            original_public_name = f"original_{i}.png"
            watermarked_public_name = f"watermarked_{i}.png"
            original_public_path = os.path.join(public_dir, original_public_name)
            watermarked_public_path = os.path.join(public_dir, watermarked_public_name)
            
            # Copy original to public directory
            shutil.copy(img_path, original_public_path)
            
            try:
                # Perform the watermarking
                embed_qr_to_image(img_path, qr_path, watermarked_path, resize_qr_if_needed=True)
                watermarked_images.append(watermarked_path)
                
                # Copy watermarked image to public directory
                shutil.copy(watermarked_path, watermarked_public_path)
                
                # Store info about this image pair with capacity info
                processed_images.append({
                    "index": i,
                    "original": f"{public_dir_name}/{original_public_name}",
                    "watermarked": f"{public_dir_name}/{watermarked_public_name}",
                    "image_size": f"{img_w}x{img_h}",
                    "capacity_ok": img_w * img_h >= qr_pixels * 4 if img_w > 0 and img_h > 0 else None
                })
                
            except Exception as e:
                print(f"[!] Gagal watermark gambar {img_path}: {str(e)}")
                # Continue with other images even if one fails

        # Replace images in the document with watermarked versions
        print(f"[*] Mengganti gambar dalam PDF dengan versi watermark")
        success = replace_images_in_pdf(pdf_path, [path for path, _, _ in valid_images], watermarked_images, output_path)

        # Clean up temporary files
        for path, _, _ in valid_images:
            if os.path.exists(path) and not os.path.basename(path).startswith("original_") and not os.path.basename(path).startswith("watermarked_"):
                os.remove(path)
        for path in watermarked_images:
            if os.path.exists(path) and not os.path.basename(path).startswith("original_") and not os.path.basename(path).startswith("watermarked_"):
                os.remove(path)

        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                print(f"[!] Warning: Tidak dapat menghapus direktori temp: {str(e)}")

        # Copy QR code to public directory for display
        qr_public_name = "watermark_qr.png"
        qr_public_path = os.path.join(public_dir, qr_public_name)
        shutil.copy(qr_path, qr_public_path)

        # Get QR code dimensions
        try:
            qr_img = Image.open(qr_path)
            qr_width, qr_height = qr_img.size
            qr_info = {
                "width": qr_width,
                "height": qr_height
            }
        except Exception as e:
            print(f"[!] Warning: Couldn't get QR dimensions: {str(e)}")
            qr_info = {"width": 21, "height": 21}  # Default QR size

        # Return success status and image info
        return {
            "success": success, 
            "processed_images": processed_images,
            "qr_image": f"{public_dir_name}/{qr_public_name}",
            "public_dir": public_dir_name,
            "qr_info": qr_info
        }
        
    except ValueError as ve:
        if str(ve) == "NO_IMAGES_FOUND":
            # Propagate the specific error
            raise ve
        print(f"[!] Error saat proses watermarking PDF: {str(ve)}")
        return {"success": False, "error": str(ve)}
    except Exception as e:
        print(f"[!] Error saat proses watermarking PDF: {str(e)}")
        return {"success": False, "error": str(e)}


def extract_watermark_from_docx(docx_path: str, output_dir: str, validate_security: bool = False, document_key: str = None) -> bool:
    """
    Extract QR watermarks from images in a docx document.

    Args:
        docx_path: Path to the .docx file
        output_dir: Directory to save extracted QR codes
        validate_security: Whether to perform security validation on extracted QR codes
        document_key: Document security key for validation (optional)

    Returns:
        bool: True if any watermarks were extracted, False otherwise
    """
    try:
        # Make sure the output directory exists
        os.makedirs(output_dir, exist_ok=True)

        # Generate document key for security validation if needed
        if validate_security and not document_key:
            print(f"[*] Generating document key for security validation...")
            document_key = security_utils.generate_document_key(docx_path)
            document_hash = security_utils.generate_document_hash(docx_path)
            print(f"[*] Document security key generated")

        # Create a unique temporary directory to store extracted images
        temp_dir_name = f"temp_extract_{uuid.uuid4().hex}"
        temp_dir = os.path.join(output_dir, temp_dir_name)
        os.makedirs(temp_dir, exist_ok=True)

        # Extract images from the document
        print(f"[*] Mengekstrak gambar dari dokumen: {docx_path}")
        extracted_images = extract_images_from_docx(docx_path, temp_dir)

        if not extracted_images:
            print("[!] Dokumen ini tidak mengandung gambar")
            # Clean up temporary directory
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except Exception as e:
                    print(f"[!] Warning: Tidak dapat menghapus direktori temp: {str(e)}")
            # Raise specific error for no images
            raise ValueError("NO_IMAGES_FOUND")

        # Try to extract QR codes from each image
        qr_found = False
        for i, img_path in enumerate(extracted_images):
            print(f"[*] Mencoba ekstraksi QR Code dari gambar {i + 1}/{len(extracted_images)}")
            qr_output_path = os.path.join(output_dir, f"extracted_qr_{i}.png")

            try:
                extract_qr_from_image(img_path, qr_output_path)
                qr_found = True

                # Try to read the QR to verify it's valid
                try:
                    qr_data = read_qr(qr_output_path)
                    if qr_data:
                        print(f"[*] QR Code berhasil diekstrak dan dibaca: {qr_data}")
                    else:
                        print(f"[!] QR Code diekstrak tetapi tidak berisi data yang valid")
                except Exception as qr_read_error:
                    print(f"[!] QR Code diekstrak tetapi tidak dapat dibaca: {str(qr_read_error)}")

            except Exception as e:
                print(f"[!] Gagal ekstraksi QR dari gambar {img_path}: {str(e)}")
                # Continue with other images even if one fails

        # Clean up temporary files
        for path in extracted_images:
            if os.path.exists(path):
                os.remove(path)

        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                print(f"[!] Warning: Tidak dapat menghapus direktori temp: {str(e)}")

        return qr_found
    except ValueError as ve:
        if str(ve) == "NO_IMAGES_FOUND":
            # Propagate the specific error
            raise ve
        print(f"[!] Error saat proses ekstraksi watermark: {str(ve)}")
        return False
    except Exception as e:
        print(f"[!] Error saat proses ekstraksi watermark: {str(e)}")
        return False


def extract_watermark_from_pdf(pdf_path: str, output_dir: str) -> bool:
    """
    Extract QR watermarks from images in a PDF document.

    Args:
        pdf_path: Path to the .pdf file
        output_dir: Directory to save extracted QR codes

    Returns:
        bool: True if any watermarks were extracted, False otherwise
    """
    try:
        # Make sure the output directory exists
        os.makedirs(output_dir, exist_ok=True)

        # Create a unique temporary directory to store extracted images
        temp_dir_name = f"temp_extract_{uuid.uuid4().hex}"
        temp_dir = os.path.join(output_dir, temp_dir_name)
        os.makedirs(temp_dir, exist_ok=True)

        # Extract images from the document
        print(f"[*] Mengekstrak gambar dari PDF: {pdf_path}")
        extracted_images = extract_images_from_pdf(pdf_path, temp_dir)

        if not extracted_images:
            print("[!] PDF ini tidak mengandung gambar")
            # Clean up temporary directory
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except Exception as e:
                    print(f"[!] Warning: Tidak dapat menghapus direktori temp: {str(e)}")
            # Raise specific error for no images
            raise ValueError("NO_IMAGES_FOUND")

        # Try to extract QR codes from each image
        qr_found = False
        for i, img_path in enumerate(extracted_images):
            print(f"[*] Mencoba ekstraksi QR Code dari gambar {i + 1}/{len(extracted_images)}")
            qr_output_path = os.path.join(output_dir, f"extracted_qr_{i}.png")

            try:
                extract_qr_from_image(img_path, qr_output_path)
                qr_found = True

                # Try to read the QR to verify it's valid
                try:
                    qr_data = read_qr(qr_output_path)
                    if qr_data:
                        print(f"[*] QR Code berhasil diekstrak dan dibaca: {qr_data}")
                    else:
                        print(f"[!] QR Code diekstrak tetapi tidak berisi data yang valid")
                except Exception as qr_read_error:
                    print(f"[!] QR Code diekstrak tetapi tidak dapat dibaca: {str(qr_read_error)}")

            except Exception as e:
                print(f"[!] Gagal ekstraksi QR dari gambar {img_path}: {str(e)}")
                # Continue with other images even if one fails

        # Clean up temporary files
        for path in extracted_images:
            if os.path.exists(path):
                os.remove(path)

        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                print(f"[!] Warning: Tidak dapat menghapus direktori temp: {str(e)}")

        return qr_found
    except ValueError as ve:
        if str(ve) == "NO_IMAGES_FOUND":
            # Propagate the specific error
            raise ve
        print(f"[!] Error saat proses ekstraksi watermark dari PDF: {str(ve)}")
        return False
    except Exception as e:
        print(f"[!] Error saat proses ekstraksi watermark dari PDF: {str(e)}")
        return False


def main():
    """Main entry point for the CLI script."""
    args = parse_arguments()

    if args.command == 'generate_qr':
        # Handle the new QR generation arguments
        version = getattr(args, 'version', None)
        error_correction = getattr(args, 'error_correction', 'M')
        box_size = getattr(args, 'box_size', 10)
        border = getattr(args, 'border', 4)
        analyze = getattr(args, 'analyze', False)
        secure = getattr(args, 'secure', False)
        document_path = getattr(args, 'document_path', None)
        
        generate_qr_code(
            data=args.data,
            output_path=args.output,
            version=version,
            error_correction=error_correction,
            box_size=box_size,
            border=border,
            analyze=analyze,
            secure=secure,
            document_path=document_path
        )

    elif args.command == 'generate_key':
        # NEW: Generate document security key
        additional_data = getattr(args, 'additional_data', '')
        save_key_path = getattr(args, 'save_key', None)
        
        generate_document_key_command(
            document_path=args.document,
            additional_data=additional_data,
            save_key_path=save_key_path
        )

    elif args.command == 'validate_security':
        # NEW: Validate QR-document security pairing
        document_key = getattr(args, 'key', None)
        detailed = getattr(args, 'detailed', False)
        
        validate_qr_document_pair(
            qr_image_path=args.qr_image,
            document_path=args.document,
            document_key=document_key,
            detailed=detailed
        )

    elif args.command == 'embed_docx':
        embed_watermark_to_docx(args.docx, args.qr, args.output)

    elif args.command == 'embed_pdf':
        embed_watermark_to_pdf(args.pdf, args.qr, args.output)

    elif args.command == 'extract_docx':
        extract_watermark_from_docx(args.docx, args.output_dir)

    elif args.command == 'extract_pdf':
        extract_watermark_from_pdf(args.pdf, args.output_dir)

    else:
        print("Perintah tidak dikenal. Gunakan --help untuk bantuan.")
        sys.exit(1)


if __name__ == "__main__":
    main()
